# -*- coding: utf-8 -*-
"""汽车租赁管理系统学术与工程实践报告生成脚本"""

import io, sys
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_PATH = '/Users/raychen/Develop/proj/CarRent/汽车租赁管理系统_学术与工程实践报告.docx'

# ── 工具函数 ──────────────────────────────────────────────
def sf(run, name='SimSun', size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold

def sp(para, *, fs=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, ls=1.25,
       sb=0, sa=0, fi=0.74, kwn=False):
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = ls
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.first_line_indent = Cm(fi) if fi else Cm(0)
    pf.keep_with_next = kwn
    for r in para.runs:
        sf(r, size=fs, bold=bold)

def ap(doc, text='', *, fs=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
       ls=1.25, sb=0, sa=0, fi=0.74, kwn=False):
    p = doc.add_paragraph()
    if text:
        sf(p.add_run(text), size=fs, bold=bold)
    sp(p, fs=fs, bold=bold, align=align, ls=ls, sb=sb, sa=sa, fi=fi, kwn=kwn)
    return p

def ah(doc, text, level):
    style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}[level]
    p = doc.add_paragraph(style=style)
    sf(p.add_run(text), size={1: 14, 2: 12, 3: 12}[level], bold=True)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(12 if level == 2 else 0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ── 表格 ──────────────────────────────────────────────────
def stf(table, wc):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa'); tblW.set(qn('w:w'), str(int(sum(wc) * 567)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement('w:tblGrid'); table._tbl.append(grid)
    for ch in list(grid): grid.remove(ch)
    for w in wc:
        c = OxmlElement('w:gridCol'); c.set(qn('w:w'), str(int(w * 567))); grid.append(c)

def ct(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fs=11):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    sf(p.add_run(text), size=fs, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def sc(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def mt(doc, hdrs, rows, wc, fs=10.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
    stf(t, wc)
    for i, h in enumerate(hdrs):
        ct(t.rows[0].cells[i], h, bold=True, fs=fs); sc(t.rows[0].cells[i], 'EDEDED')
    for ri, row in enumerate(rows, start=1):
        for i, v in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.LEFT if i == len(row) - 1 and len(wc) > 3 else WD_ALIGN_PARAGRAPH.CENTER
            ct(t.rows[ri].cells[i], str(v), align=al, fs=fs)

# ── 域代码 ──────────────────────────────────────────────
def af(para, fc):
    r = para.add_run()
    for t, tv in [('begin', None), ('instr', fc), ('separate', None), ('text', '1'), ('end', None)]:
        if t == 'text':
            e = OxmlElement('w:t'); e.text = tv
        elif t == 'instr':
            e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = fc
        else:
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), t)
        r._r.append(e)

def pnf(sec):
    sec.footer.is_linked_to_previous = False
    p = sec.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    af(p, ' PAGE ')
    for r in p.runs: sf(r, size=9)

def ssl(sec, spn=None):
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
    sec.header_distance = Cm(1.2); sec.footer_distance = Cm(1.2)
    if spn:
        pgnt = sec._sectPr.find(qn('w:pgNumType'))
        if pgnt is None:
            pgnt = OxmlElement('w:pgNumType'); sec._sectPr.append(pgnt)
        pgnt.set(qn('w:start'), str(spn))

def ssty(doc):
    n = doc.styles['Normal']
    n.font.name = 'SimSun'; n._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    n.font.size = Pt(12)
    n.paragraph_format.line_spacing = 1.25
    for sn, sz, b, al in [('Title', 18, True, WD_ALIGN_PARAGRAPH.CENTER),
                           ('Heading 1', 14, True, WD_ALIGN_PARAGRAPH.LEFT),
                           ('Heading 2', 12, True, WD_ALIGN_PARAGRAPH.LEFT),
                           ('Heading 3', 12, True, WD_ALIGN_PARAGRAPH.LEFT)]:
        s = doc.styles[sn]; s.font.name = 'SimSun'
        s._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        s.font.size = Pt(sz); s.font.bold = b
        s.paragraph_format.line_spacing = 1.25
        s.paragraph_format.space_after = Pt(6 if 'Heading' in sn else 0)
        s.paragraph_format.keep_with_next = True

def add_code(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.0; pf.first_line_indent = Cm(0)
        sf(p.add_run(line), name='Courier New', size=8)
    ap(doc, '', sa=6, fi=0)


# ── 子节构建 ──────────────────────────────────────────────
def add_subsect(doc, num, title, overview, code, usage, difficulty):
    ah(doc, '2.2.%d %s' % (num, title), 3)
    ap(doc, '■ 功能描述', bold=True, fs=12, sb=6, sa=3, fi=0)
    ap(doc, overview, sa=6)
    ap(doc, '■ 主要源代码（含注释）', bold=True, fs=12, sb=6, sa=3, fi=0)
    add_code(doc, code)
    ap(doc, '■ 使用说明', bold=True, fs=12, sb=6, sa=3, fi=0)
    ap(doc, usage, sa=6)
    ap(doc, '■ 技术难点', bold=True, fs=12, sb=6, sa=3, fi=0)
    ap(doc, difficulty, sa=6)


# ── 主构建函数 ────────────────────────────────────────────
def build():
    doc = Document()
    ssty(doc)
    for s in doc.sections: ssl(s)

    # ===================== 封面 ======================================
    ap(doc, '2025-2026年春季学期', fs=14, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0)
    doc.add_paragraph()
    ap(doc, '学术与工程实践(计算机)', fs=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0)
    ap(doc, '课设报告', fs=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24, fi=0)
    for l, r in [('课设题目：', '汽车租赁管理系统'), ('指导老师：', '孙晶'),
                 ('班    级：', '________________'), ('学    号：', '________________'),
                 ('姓    名：', '________________')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.sb = Pt(0); p.paragraph_format.sa = Pt(8)
        p.paragraph_format.fi = Cm(0)
        sf(p.add_run(l), size=12); sf(p.add_run(r), size=12)
    doc.add_paragraph(); doc.add_paragraph()
    ap(doc, '总成绩：        验机成绩：      报告成绩：', fs=12, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0)
    ap(doc, '评语：', fs=12, align=WD_ALIGN_PARAGRAPH.LEFT, fi=0)
    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
    ap(doc, '指导教师签字：             日期：', fs=12, align=WD_ALIGN_PARAGRAPH.RIGHT, fi=0)

    # ===================== 目录节 ======================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    ssl(doc.sections[1])
    ap(doc, '目录', fs=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=18, fi=0)
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    af(p, 'TOC \\o "1-3" \\h \\z \\u')

    # ===================== 正文节 ======================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    ssl(doc.sections[2], spn=1)
    pnf(doc.sections[2])

    # ── 第 1 章 ──────────────────────────────────────
    ah(doc, '1 课题主要内容介绍', 1)
    ap(doc, '汽车租赁管理系统是面向中小型汽车租赁公司日常业务的信息化管理工具。系统基于 C++ 语言开发，采用命令行交互方式，结合单向链表组织数据和二进制文件持久化存储，实现了车辆信息管理、租车用户信息管理、租车业务办理、退车业务办理、信息查询、统计汇总、系统设置等七大核心模块。', sa=6)
    ap(doc, '本课题通过将结构体、指针、链表、字符串处理、日期处理、文件操作和交互式界面等知识点融合在一个完整项目中，旨在训练学生对中小型管理系统的需求分析、模块拆分和程序组织能力。', sa=6)

    ah(doc, '1.1 课题概述', 2)
    ap(doc, '实训目的：通过创建命令行模式的汽车租赁管理系统，掌握结构体类型的定义与引用方法、文件打开/关闭/读写/定位操作、通过动态链表对数据的增删改查操作，同时培养 C/C++ 语言程序编程规范。', sa=6)
    ap(doc, '项目需求：系统为汽车租赁公司提供车辆信息管理、租用者信息管理、租赁业务信息管理等功能，支持按不同条件查询统计车辆租赁情况和公司盈利情况，为公司经营决策提供依据。租用者需提供驾照、姓名等自然信息以及租用车型、起始时间等信息；退车时需办理缴费，系统自动计算费用并打印票据。', sa=6)
    ap(doc, '项目要求：系统需支持车辆信息管理（增删改）、用户信息管理（增删改）、租车与退车业务办理、票据生成、信息查询（按类别、按日期区间、模糊搜索）、统计汇总（按品牌/类型/颜色/性别/驾龄/月份分组）以及系统登录密码保护。数据以二进制文件持久化，运行时可从文件加载并构建链表。', sa=6)

    ah(doc, '1.2 运行环境说明', 2)
    mt(doc,
       ['项目', '说明', '项目', '说明'],
       [['开发语言', 'C++11', '编译器', 'g++ / Clang / MSVC'],
        ['构建方式', 'Make / CMake', '运行方式', '命令行交互'],
        ['存储方式', '二进制文件 + 文本日志', '编码', 'UTF-8'],
        ['适用平台', 'macOS / Linux / Windows', '开发环境', 'CLion / VS Code']],
       [3.0, 5.4, 3.0, 5.4])
    ap(doc, '编译入口采用单翻译单元模式，main.cpp 通过 #include 直接引入核心实现文件和各菜单模块，最终只需编译 main.cpp 即可生成可执行文件。数据文件包括 vehicle.dat、renter.dat、rent.dat、pwd.dat 和 log.txt，均由系统自动管理。', sa=6)

    doc.add_page_break()

    # ── 第 2 章 ──────────────────────────────────────
    ah(doc, '2 系统设计与实现', 1)
    ap(doc, '本章从设计原则、功能划分、数据结构设计和七个功能模块的实现细节四个方面展开。系统整体采用"公共核心 + 菜单拆分"的结构，核心文件 car_rent.cpp 负责定义数据模型、输入工具、日期工具、持久化、统计函数和主菜单，七个菜单文件各自负责对应模块的交互流程。', sa=6)

    ah(doc, '2.1 系统设计', 2)
    ah(doc, '2.1.1 设计原则与设计思路', 3)
    ap(doc, '设计原则一：数据一致性。车辆、用户和租赁记录之间存在明确关联，每次租车或退车操作都必须同步更新多个实体，不能只修改局部字段。系统通过统一的查找函数和保存函数保证链路一致，通过状态枚举（VehicleStatus / RentStatus）进行业务流转。', sa=6)
    ap(doc, '设计原则二：输入可靠性。系统使用 fgets() + strtol() / strtod() 替代 scanf / atoi / atof，对整数、浮点数、日期和字符均进行合法性校验和范围判断，能够区分空输入、非数字输入和超范围输入。', sa=6)
    ap(doc, '设计原则三：持久化简洁可控。系统不使用外部数据库，而是将链表数据一次性写入二进制文件，启动时重新读入并恢复链表关系。数据文件格式为 count + 记录块，读者可根据需要直接解析。', sa=6)
    ap(doc, '设计思路：项目将业务拆为七个菜单，主菜单只负责调度，各子菜单负责收集输入并调用核心函数。控制流采用"输入 -> 校验 -> 核心操作 -> 持久化保存 -> 结果反馈"的标准化流程。', sa=6)

    ah(doc, '2.1.2 功能设计', 3)
    mt(doc,
       ['模块', '名称', '主要职责'],
       [['菜单 1', '车辆信息管理', '添加/删除/修改/显示车辆，车牌号重复性校验'],
        ['菜单 2', '用户信息管理', '添加/删除/修改/显示用户，驾照号、身份证号唯一性校验'],
        ['菜单 3', '租车业务办理', '列出所有可用车辆，录入日期，生成租车记录，打印租车票据'],
        ['菜单 4', '退车业务办理', '输入归还日期，自动计算天数、费用和退款，打印退车收据'],
        ['菜单 5', '信息查询', '按品牌/类型/颜色/状态/日期区间筛选；跨实体模糊关键词搜索'],
        ['菜单 6', '统计汇总与输出', '系统概览、车辆/用户/租车统计（ASCII柱状图）、全部信息输出'],
        ['菜单 7', '系统设置', '修改密码、导出报表、查看操作日志、批量清空测试数据']],
       [2.0, 3.6, 11.2], fs=10.0)
    ap(doc, '', sa=6, fi=0)

    ah(doc, '2.1.3 数据结构设计', 3)
    ap(doc, '系统定义了三类核心结构体：Vehicle（车辆）、Renter（用户）、RentRecord（租车记录），每类结构体均通过对应的链表节点进行动态组织。', sa=6)

    add_code(doc, '''Vehicle（车辆结构体）
    id: int         自动递增编号
    plateNo[16]:    车牌号（唯一约束）
    brand[32]:      品牌
    type[32]:       类型
    color[16]:      颜色
    purchaseDate[16]: 购买日期（YYYY-MM-DD）
    status: int     状态：0-可租, 1-已租, 2-维修
    dailyRate: double 日租金
    mileage: double 里程
    insurance[64]:  保险信息''')

    add_code(doc, '''Renter（用户结构体）
    id: int         自动递增编号
    name[32]:       姓名
    age: int        年龄
    gender: char    性别（M / F）
    phone[20]:      联系电话
    licenseNo[20]:  驾照号码（唯一约束）
    idCard[20]:     身份证号（唯一约束）
    drivingYears: int 驾龄
    rentCount: int  累计租车次数''')

    add_code(doc, '''RentRecord（租车记录结构体）
    id: int         自动递增编号
    vehicleId: int  关联车辆编号
    renterId: int   关联用户编号
    rentDate[16]:   租车日期
    expectedReturnDate[16]: 预计归还日期
    actualReturnDate[16]:   实际归还日期
    deposit: double 押金（日租金 x 3）
    dailyRate: double 快照日租金
    totalFee: double 实际总费用
    status: int     0-租用中, 1-已归还
    vehicleBrand[32]: 快照：车辆品牌
    vehiclePlate[16]: 快照：车牌号
    renterName[32]:   快照：用户姓名
    renterLicense[20]:  快照：驾照号''')

    ap(doc, '链表节点定义：VehicleNode 包含 Vehicle data 和 VehicleNode* next；RenterNode 和 RentNode 同理。三组全局头指针 vehicleHead、renterHead、rentHead 分别指向各自链表。这种组织方式在插入和删除时有 O(n) 复杂度，但在课程设计规模下足够使用，且便于遍历、统计和导出。', sa=6)

    mt(doc,
       ['数据文件', '存储格式', '说明'],
       [['vehicle.dat', '[count:4B] [Vehicle x N]', '车辆链表持久化文件'],
        ['renter.dat', '[count:4B] [Renter x N]', '用户链表持久化文件'],
        ['rent.dat', '[count:4B] [RentRecord x N]', '租赁记录持久化文件'],
        ['pwd.dat', '[username:32B][pwd_enc:64B]', '登录凭据（XOR加密）'],
        ['log.txt', 'YYYY-MM-DD HH:MM:SS 操作描述', '操作日志文件']],
       [4.0, 6.0, 6.8], fs=10.0)
    ap(doc, '', sa=6, fi=0)

    ah(doc, '2.2 系统实现', 2)

    # ---- 2.2.1 车辆信息管理 ----
    add_subsect(doc, 1, '车辆信息管理功能',
        '车辆管理模块负责车辆基本信息的录入、维护和展示。添加车辆时，系统依次提示输入车牌号、品牌、类型、颜色、购买日期、保险信息、日租金和里程。输入阶段进行三重校验：车牌号重复检查（遍历 vehicleHead 链表做 strcmp）、数值范围校验（日租金与里程非负）、以及购买日期格式校验（isValidDate 含闰年识别）。所有新增车辆默认状态为可租，并分配新的自增编号，最后新建 VehicleNode 追加至链表末尾并写回 vehicle.dat。\n\n删除操作通过车辆 ID 在链表中定位目标节点，记录前驱指针后执行 prev->next = cur->next 并 delete 释放内存。修改操作保持原 ID 不变，仅替换其余字段，保证后续租赁记录仍能正确关联。显示所有车辆时使用统一的 printf 表头与按列对齐方式输出，每 20 行分页暂停。',
        '''void menu1Vehicle() {
    while (true) {
        int c = inputInt("请选择: ", 0, 4);
        if (c == 0) return;
        if (c == 4) { printAllVehicles(); pauseScreen(); continue; }
        if (c == 1) {
            Vehicle v{};  initVehicle(v);       // memset清零
            while (true) {
                inputLine("车牌号: ", v.plateNo, MAX_PLATE_LEN);
                if (!isPlateNoDuplicate(v.plateNo)) break;
                printf("车牌号已存在\\n");
            }
            inputLine("品牌: ", v.brand, MAX_BRAND_LEN);
            v.dailyRate = inputDouble("日租金: ", 0.0, 999999.0);
            v.mileage = inputDouble("里程: ", 0.0, 9999999.0);
            addVehicle(v);   // new节点->赋值->链到末尾->保存
            printf("已添加\\n"); pauseScreen();
        } else if (c == 2) {
            printAllVehicles();
            printf(deleteVehicle(inputInt("ID: ",1,9999))?"已删除":"没找到");
            pauseScreen();
        } else if (c == 3) {
            Vehicle* p = findVehicle(inputInt("ID: ",1,9999));
            if (!p) { printf("没找到\\n"); pauseScreen(); continue; }
            Vehicle v = *p;
            // ... 逐字段重新输入 ...
            printf(modifyVehicle(v.id, v) ? "已修改\\n" : "失败\\n");
            pauseScreen();
        }
    }
}''',
        '使用者录入数据时必须按照系统提示逐字段输入。车牌号不允许重复，同一车牌只能录入一次。日租金和里程必须为非负数值。购买日期需按 YYYY-MM-DD 格式输入，系统会校验闰年和每月天数。删除和修改操作通过 ID 定位目标记录，操作完成后数据立即写入文件；若 ID 不存在则提示"没找到"。',
        '技术难点一：车牌号重复检查需要在每次新增时遍历整个链表，对于已有大量记录的场景有性能损耗，但在本项目中规模可控。\n技术难点二：日期合法性校验涉及闰年逻辑和不同月份的天数差异，实现时需在二月判断是否为闰年以决定 28 天还是 29 天。\n技术难点三：链表删除操作必须同时维护前驱指针，避免链表断裂导致内存泄漏或数据丢失。')

    # ---- 2.2.2 用户信息管理 ----
    add_subsect(doc, 2, '用户信息管理功能',
        '用户信息管理功能与车辆管理类似，但更强调身份信息的唯一性。新增用户时需要对驾照号和身份证号分别进行链表的唯一性检查，若有重复则要求重新输入。用户信息包括姓名、年龄、性别、电话、驾照号、身份证号、驾龄和租车次数（初始化为 0）。性别采用 M/F 双值约束，在输入后统一转为大写字母（M / F），保证后续查询和统计的一致性。年龄限制在 1 到 120 之间，驾龄限制在 0 到 60 年之间。',
        '''void menu2Renter() {
    while (true) {
        int c = inputInt("请选择: ", 0, 4);
        if (c == 0) return;
        if (c == 4) { printAllRenters(); pauseScreen(); continue; }
        if (c == 1) {
            Renter r;  initRenter(r);           // memset清零
            inputLine("姓名: ", r.name, MAX_NAME_LEN);
            r.age = inputInt("年龄: ", 1, 120);
            r.gender = inputChar("性别(M/F): ", "MmFf");
            if (r.gender == 'm') r.gender = 'M'; // 统一大写
            if (r.gender == 'f') r.gender = 'F';
            while (true) {
                inputLine("驾照号: ", r.licenseNo, MAX_LICENSE_LEN);
                if (!isLicenseNoDuplicate(r.licenseNo)) break;
                printf("驾照号已存在\\n");
            }
            while (true) {
                inputLine("身份证号: ", r.idCard, MAX_IDCARD_LEN);
                if (!isIdCardDuplicate(r.idCard)) break;
                printf("身份证号已存在\\n");
            }
            r.drivingYears = inputInt("驾龄: ", 0, 60);
            addRenter(r);   // 分配ID->new节点->链到末尾->保存
        }
    }
}''',
        '使用者需按照系统提示逐字段输入用户信息。驾照号和身份证号在系统中均不允许重复，若录入已被其他用户使用的证件号，系统会提示重新输入。性别输入 M 或 F（不区分大小写），由系统自动统一为大写。年龄和驾龄有合法范围限制，超出范围会被拒绝。',
        '技术难点一：驾照号与身份证号是不同于车牌号的独立查重逻辑，需要单独编写两个查重函数，分别遍历 renterHead 链表进行比较。\n技术难点二：性别输入的大小写处理。用户可能输入 m/M 或 f/F，系统统一转为大写，避免后续查询和统计不一致。\n技术难点三：用户修改操作保留原 ID 和租车次数，修改后不会丢失该用户的历史租赁记录关联。')

    doc.add_page_break()

    # ---- 2.2.3 租车业务办理 ----
    add_subsect(doc, 3, '租车业务办理功能',
        '租车办理是最核心的业务链路。进入该模块后，系统先通过 countAvailableVehicles() 统计可用车辆数量，若无可用车辆，则直接提示"没有剩余车辆"并结束。若有可用车辆，则遍历链表列出所有状态为可租的车辆供用户选择。\n\n用户输入车辆 ID、用户 ID、租车日期和预计归还日期后，系统对日期格式进行严格校验（isValidDate 含闰年识别与每月天数验证）。addRentRecord() 函数验证两 ID 的有效性及车辆状态，通过后在堆上 new RentNode，将记录数据（含车辆品牌/车牌号/用户姓名/驾照号的快照字段）赋值并链入链表末尾，再将车辆状态改为已租、用户租车次数加 1，最后统一写入三个数据文件。押金按照日租金 x 3 自动计算。成功后打印格式化的"汽车租赁公司租车票据"。',
        '''void menu3Rent() {
    while (true) {
        int c = inputInt("请选择: ", 0, 2);
        if (c == 0) return;
        if (c == 2) { printAllRents(); pauseScreen(); continue; }
        int avail = countAvailableVehicles();   // 统计可用车数量
        if (avail == 0) {
            printf("没有剩余车辆。\\n"); pauseScreen(); continue;
        }
        VehicleNode* cur = vehicleHead;
        while (cur) {
            if (cur->data.status == STATUS_AVAILABLE)
                printVehicleRow(cur->data);     // 只输出可租车辆
            cur = cur->next;
        }
        int vehicleId = inputInt("车辆ID: ", 1, 9999);
        int renterId = inputInt("用户ID: ", 1, 9999);
        char rd[16], ed[16];
        inputLine("租车日期: ", rd, 16);
        inputLine("预计归还日期: ", ed, 16);
        if (!isValidDate(rd) || !isValidDate(ed)) continue;
        if (!addRentRecord(vehicleId, renterId, rd, ed)) {
            printf("办理失败\\n"); pauseScreen(); continue;
        }
        // 找到刚创建的记录打印票据
        RentNode* rn = rentHead;
        while (rn) {
            if (rn->data.vehicleId == vehicleId &&
                rn->data.renterId == renterId) break;
            rn = rn->next;
        }
        if (rn) printRentTicket(rn->data);
        pauseScreen();
    }
}''',
        '1. 进入租车办理后，系统自动列出全部可租车辆供选择。\n2. 输入车辆 ID 和用户 ID，系统校验两者有效性。\n3. 租车日期和预计归还日期需按 YYYY-MM-DD 格式输入。\n4. 办理成功后，系统同时更新车辆状态和用户租车次数，并打印格式化票据。\n5. 若没有可用车辆，系统直接提示，避免用户继续操作。',
        '技术难点一：多实体联动更新。租车操作需要同时处理 Vehicle、Renter 和 RentRecord 三类对象，addRentRecord() 使用指针直接修改对象，在保证原子性的同时避免额外数据复制。\n技术难点二：快照字段设计。租车记录中冗余保存了品牌/车牌/姓名/驾照号，即使后续信息被修改，历史票据仍保留原始信息。\n技术难点三：日期格式校验必须在办理前完成。isValidDate() 逐段解析年/月/日并做闰年判断，拒绝不符合格式的输入。')

    # ---- 2.2.4 退车业务办理 ----
    add_subsect(doc, 4, '退车业务办理功能',
        '退车办理以租车记录 ID 为入口。系统先通过 findRent 在 rentHead 链表中定位对应记录，若记录不存在或状态已为已归还，则提示办理失败。\n\n费用计算公式：总费用 = 租车天数 x 日租金，退款 = 押金 - 总费用。其中租车天数通过 calcDateDiff() 将租车日期和归还日期分别转换为自 2000-01-01 起的天数后求差值，最少按 1 天计费（防止同日租还产生 0 天费用）。若退款小于 0 则截断为 0。\n\n完成后，租赁记录状态改为已归还，对应车辆恢复为可租，数据写入文件。最后打印格式化的"汽车租赁公司租车费用收据"。',
        '''void menu4Return() {
    while (true) {
        int c = inputInt("请选择: ", 0, 3);
        if (c == 0) return;
        int rentId = inputInt("租车记录ID: ", 1, 9999);
        RentRecord* r = findRent(rentId);
        if (!r) { printf("没找到\\n"); pauseScreen(); continue; }
        if (c == 3) {
            printRentTicket(*r); pauseScreen(); continue;
        }
        char returnDate[MAX_DATE_LEN];
        inputLine("实际归还日期: ", returnDate, MAX_DATE_LEN);
        double totalFee = 0, refund = 0;
        if (!returnRent(rentId, returnDate, totalFee, refund)) {
            printf("办理失败\\n"); pauseScreen(); continue;
        }
        RentRecord* rec = findRent(rentId);
        int days = calcDateDiff(rec->rentDate, returnDate);
        if (days < 1) days = 1;                // 最少1天
        printReturnReceipt(*rec, days, totalFee, refund);
        pauseScreen();
    }
}''',
        '1. 输入要退车的租车记录 ID。\n2. 系统校验该记录是否存在及是否已归还。\n3. 选项 3 可单独查看某条记录的详细信息。\n4. 退车日期需按规范格式输入，系统自动计算天数和费用。\n5. 退车完成后打印收据，展示完整费用构成。',
        '技术难点一：日期差计算需处理跨月、跨年、闰年等复杂情况。使用 dateToDays() 归一化后求差，同时规定最少 1 天计费。\n技术难点二：退车后车辆状态必须恢复为可租，确保后续用户能再次租赁。\n技术难点三：退款若为负值（超期导致费用大于押金），系统截断为 0。')

    doc.add_page_break()

    # ---- 2.2.5 信息查询 ----
    add_subsect(doc, 5, '信息查询功能',
        '查询模块提供两级菜单结构。第一级让用户选择"分类查询"或"模糊搜索"。分类查询支持车辆按品牌/类型/颜色/状态、用户按姓名/驾照号/性别、租车记录按状态/日期区间进行筛选。所有查询基于链表遍历实现，命中条件后调用统一的打印函数输出。\n\n模糊搜索采用 strstr() 子串匹配，同时搜索车辆的品牌+车牌号和用户的姓名+驾照号，适合在记不清完整编号时快速定位。日期区间查询将起止日期通过 dateToDays() 转化为天数后做范围比较。性别查询支持大小写不敏感，通过 cur->data.gender == key[0] || cur->data.gender == (key[0] - 32) 实现。',
        '''void menu5Query() {
    // 1. 模糊搜索：strstr 跨实体匹配
    char key[64]; inputLine("关键词: ", key, sizeof(key));
    VehicleNode* vc = vehicleHead;
    while (vc) {
        if (strstr(vc->data.brand, key) || strstr(vc->data.plateNo, key))
            printVehicleRow(vc->data);
        vc = vc->next;
    }
    // 用户侧同理：name + licenseNo
    RenterNode* rc = renterHead;
    while (rc) {
        if (strstr(rc->data.name, key) || strstr(rc->data.licenseNo, key))
            printRenterRow(rc->data);
        rc = rc->next;
    }
    // 2. 分类查询：按品牌/类型/颜色/状态等
    VehicleNode* cur = vehicleHead;
    while (cur) {
        bool match = false;
        if (sc == 1) match = strstr(cur->data.brand, key);
        else if (sc == 2) match = strstr(cur->data.type, key);
        else if (sc == 3) match = strstr(cur->data.color, key);
        else if (sc == 4) match = (cur->data.status == st);
        if (match) printVehicleRow(cur->data);
        cur = cur->next;
    }
    // 3. 日期范围查询：转为天数后数值比较
    int start = dateToDays(d1), end = dateToDays(d2);
    while (cur) {
        int rd = dateToDays(cur->data.rentDate);
        if (rd >= start && rd <= end) printRentRow(cur->data);
        cur = cur->next;
    }
}''',
        '1. 选择"分类查询"进入按维度筛选模式，选择"模糊搜索"进入关键词跨实体搜索模式。\n2. 分类查询：车辆支持按品牌/类型/颜色（子串匹配）和状态查询。用户支持按姓名/驾照号/性别查询。租车记录支持按状态和日期范围查询。\n3. 模糊搜索同时匹配车辆与用户，输入一个关键词即可看到两方的匹配结果。\n4. 所有查询结果以表格形式输出。',
        '技术难点一：分类查询条件分支较多，通过统一 bool match 标志和逐项判断实现。\n技术难点二：日期区间查询需将字符串日期转为可比较的数值。通过 dateToDays 归一化，避免字符串比较带来的月份/日期错位问题。\n技术难点三：模糊搜索的跨实体匹配——一个关键词同时搜索车辆和用户两类实体，用 strstr 统一处理。')

    # ---- 2.2.6 统计汇总 ----
    add_subsect(doc, 6, '统计汇总功能',
        '统计模块提供系统概览、车辆统计、用户统计、租车统计和全部信息输出五项能力。系统概览部分输出总数统计并以 ASCII 柱状图展示车辆状态分布与品牌分布。分组统计通过专门的计数函数实现：countVehiclesByBrand() 遍历车辆链表，用 strcmp 去重后各自计数；countRentersByDrivingYears() 按驾驶年限分为 0-3 年、4-10 年、10 年以上三段；countRentsByMonth() 从租车日期的第 6-7 位提取月份并计入对应槽位。\n\nASCII 柱状图由 drawBarChart() 统一绘制：先遍历找出最大值，然后将各值乘以 30 除以最大值确定 # 数量，最后输出标签、数量和柱状条。\n\n输出全部信息功能依次调用 printAllVehicles()、printAllRenters()、printAllRents()，每 20 行分页暂停。',
        '''// 车辆按品牌分组统计
int countVehiclesByBrand(char labels[][32], int values[], int max) {
    int n = 0;
    VehicleNode* cur = vehicleHead;
    while (cur) {
        int found = -1;
        for (int i = 0; i < n; i++)     // 查找该品牌是否已出现过
            if (strcmp(labels[i], cur->data.brand) == 0)
                { found = i; break; }
        if (found >= 0) values[found]++;      // 已出现：计数+1
        else if (n < max) {                   // 新品牌：添加标签
            strncpy(labels[n], cur->data.brand, 31);
            values[n] = 1; n++;
        }
        cur = cur->next;
    }
    return n;
}

// ASCII 柱状图
void drawBarChart(char labels[][32], int values[], int n) {
    int mx = 0;
    for (int i = 0; i < n; i++)
        if (values[i] > mx) mx = values[i];      // 1. 找最大值
    for (int i = 0; i < n; i++) {
        printf("  %-10s [%3d] ", labels[i], values[i]);
        int bar = mx > 0 ? values[i] * 30 / mx : 0;  // 3. 按比例
        for (int j = 0; j < bar; j++) printf("#");    // 4. 绘制
        printf("\\n");
    }
}

// 驾龄分段统计
void countRentersByDrivingYears(int values[3]) {
    values[0] = values[1] = values[2] = 0;
    RenterNode* cur = renterHead;
    while (cur) {
        int y = cur->data.drivingYears;
        if (y <= 3) values[0]++;         // 0-3 年
        else if (y <= 10) values[1]++;   // 4-10 年
        else values[2]++;                // 10 年以上
        cur = cur->next;
    }
}''',
        '1. 系统概览：自动输出总数统计 + 车辆状态柱状图 + 品牌分布柱状图。\n2. 车辆统计：选择按品牌/类型/颜色分组，每组以 ASCII 柱状图展示。\n3. 用户统计：选择按性别（M/F）或驾龄分段（0-3年/4-10年/10年+）统计。\n4. 租车统计：按 1-12 月统计每月租车数量。\n5. 输出全部信息：一次性展示三张完整表格，每 20 行暂停。',
        '技术难点一：分组统计的去重逻辑需要手动维护 labels 数组，遍历时先线性检索已有标签再决定是否新增。\n技术难点二：ASCII 柱状图需要处理"所有值为 0"的特殊情况，在 mx > 0 时才做除法。\n技术难点三：驾龄分段统计需要将连续数值映射到离散区间，注意边界条件（10 年归入第二段而非第三段）。')

    doc.add_page_break()

    # ---- 2.2.7 系统设置 ----
    add_subsect(doc, 7, '系统设置与登录功能',
        '系统设置模块提供密码修改、报表导出、日志查看和数据清空功能。\n\n登录功能：首次启动时密码文件不存在或为空，系统自动进入初始化流程，引导用户设置用户名和密码（密码长度 6-16 位），两次输入一致后通过 XOR 加密写入 pwd.dat。之后每次启动均需通过用户名和密码验证，最多三次尝试。\n\n密码修改：验证旧密码 -> 新密码两次输入一致 -> XOR 加密后覆写 pwd.dat。\n\n数据清空：确认后遍历三条链表逐个 delete 所有节点，计数器归零，保存空文件覆盖。',
        '''// XOR 加密
void xorPassword(const char* plain, char* out) {
    int len = static_cast<int>(strlen(plain));
    for (int i = 0; i < 64; i++)
        out[i] = i < len ? static_cast<char>(plain[i] ^ 0x5A) : 0;
}

// 登录（用户名+密码双重验证）
bool login() {
    if (!passwordReady) {
        // 首次使用：设置用户名和密码
        inputLine("设置用户名: ", usernameStore, 32);
        char p1[64], p2[64];
        while (true) {
            inputLine("设置密码: ", p1, 64);
            inputLine("确认密码: ", p2, 64);
            if (strcmp(p1, p2) == 0 && setNewPassword(p1)) break;
        }
        savePassword();
        return true;
    }
    for (int i = 0; i < 3; i++) {
        char u[64], p[64];
        inputLine("用户名: ", u, 64);
        inputLine("密码: ", p, 64);
        if (strcmp(u, usernameStore) == 0 && checkPassword(p))
            return true;
    }
    return false;
}

// 清空数据
while (vehicleHead) { VehicleNode* t = vehicleHead; vehicleHead = vehicleHead->next; delete t; }
while (renterHead)   { RenterNode*   t = renterHead;   renterHead = renterHead->next;   delete t; }
while (rentHead)     { RentNode*     t = rentHead;     rentHead = rentHead->next;       delete t; }
vehicleCount = renterCount = rentCount = 0;
saveAllData();''',
        '1. 首次运行设置用户名和密码（密码 6-16 位，两次需一致）。\n2. 后续每次启动需输入用户名+密码，错误超过 3 次自动退出。\n3. 修改密码需先验证旧密码。\n4. 清空数据会释放所有链表节点，操作前有二次确认。\n5. 报表导出和日志查看为只读操作。',
        '技术难点一：XOR 加密使用 0x5A 作为密钥进行逐字节异或，文件以密文存储。\n技术难点二：memcmp 比对加密结果时必须使用常量比较（整个 64 字节数组）。\n技术难点三：数据清空需要先逐个释放链表节点再清空计数器，顺序不能颠倒，否则导致内存泄漏。')

    doc.add_page_break()

    # ── 第 3 章 ──────────────────────────────────────
    ah(doc, '3 系统测试和结果分析', 1)
    ap(doc, '本章使用代表性测试数据验证系统的核心功能，包括登录安全、重复性检测、租车退车计费、日期校验和查询统计。测试在 macOS 终端环境下使用 g++ -std=c++11 编译产生。', sa=6)

    ah(doc, '3.1 登录安全测试', 2)
    mt(doc,
       ['测试项', '输入', '预期结果', '实际结果', '分析'],
       [['首次初始化', '用户名=admin,密码=123456', '提示设置成功，pwd.dat为密文', '成功，pwd.dat生成密文', '首次启动自动初始化，不支持跳过'],
        ['正常登录', '用户名=admin,密码=123456', '进入主菜单', '成功进入，显示标题栏和时钟', 'XOR解密比对正确'],
        ['密码错误', '用户名=admin,密码=000000', '提示错误，剩余2次', '提示用户名或密码错误', '错误计数正常递减'],
        ['三次错误退出', '连续三次错误密码', '程序退出', '第三次后login返回false', '防止暴力破解']],
       [2.4, 3.6, 3.6, 3.6, 3.6], fs=9.5)
    ap(doc, '登录模块能区分首次初始化和正常登录两种状态。密码写入 pwd.dat 前经过 XOR 0x5A 加密，直接用 cat 命令查看时仅显示乱码。', sa=6)

    ah(doc, '3.2 输入校验与重复检查测试', 2)
    mt(doc,
       ['测试项', '输入', '预期结果', '实际结果', '分析'],
       [['车牌号重复', '添加已有车牌', '提示车牌号已存在，重新输入', '符合预期', 'isPlateNoDuplicate遍历链表返回true'],
        ['驾照号重复', '添加已有驾照号', '提示驾照号已存在', '符合预期', 'isLicenseNoDuplicate返回true'],
        ['身份证号重复', '添加已有身份证号', '提示身份证号已存在', '符合预期', 'isIdCardDuplicate返回true'],
        ['空输入', '整数输入直接回车', '提示输入不能为空', '符合预期', 'line[0]==chr(0)判断生效'],
        ['非法字符', '整数输入hello', '提示请输入有效的整数', '符合预期', 'strtol检测*end非ch(0)后拒绝'],
        ['日期格式错误', '输入2026-02-30', '提示日期格式错误', '符合预期', 'isValidDate验证2月最多29天']],
       [2.4, 3.2, 4.0, 3.6, 3.6], fs=9.5)
    ap(doc, '输入模块的 inputInt/inputDouble/inputChar 对所有交互输入进行规范化处理。空输入、非数字输入、超范围输入分别对应三种提示信息。', sa=6)

    ah(doc, '3.3 租车与退车计费测试', 2)
    mt(doc,
       ['测试项', '输入', '预期结果', '实际结果', '分析'],
       [['正常租车', '车1,人1,租07-10,还07-13', '票据显示押金=日租金x3', '符合预期', 'addRentRecord正确处理'],
        ['无可用车', '车辆全部已租时选3', '提示没有剩余车辆', '符合预期', 'countAvailableVehicles返回0'],
        ['正常退车', '记录1,归还07-13', '3天x日租金=总费用,退款=押金-总费用', '符合预期', 'calcDateDiff计算为3天'],
        ['同日退车', '租车与归还日期相同', '最少按1天计费', '费用=1x日租金', 'days<1则days=1逻辑生效'],
        ['超期退车', '实际天数远超预计', '费用可能超押金,退款截断为0', '退款为0.00', 'refund<0则refund=0生效']],
       [2.4, 4.0, 4.0, 3.6, 2.8], fs=9.5)
    ap(doc, '租车退车计费逻辑经过多组边界测试全部正确。同日租退场景按最低 1 天计费；超期使用场景通过退款截断保护公司财务安全。', sa=6)

    ah(doc, '3.4 查询与统计测试', 2)
    mt(doc,
       ['测试项', '输入', '预期结果', '实际结果', '分析'],
       [['按品牌查车', '关键词=丰田', '输出含丰田品牌的车辆', '符合预期，1条命中', 'strstr子串匹配'],
        ['模糊搜索', '关键词=A', '车辆品牌/车牌+用户名/驾照命中', '2条车辆+1条用户', '跨实体搜索有效'],
        ['日期区间查询', '起始=07-01,结束=07-31', '仅输出该月租车记录', '符合预期', 'dateToDays归一化后范围筛选'],
        ['性别统计', '用户统计->性别', '柱状图输出M和F数量', '图形正常，数值正确', 'countRentersByGender遍历'],
        ['驾龄分段统计', '0-3/4-10/10+', '三段分别计数', '与原始数据一致', 'if-else分段正确']],
       [2.4, 4.2, 3.6, 3.6, 3.0], fs=9.5)
    ap(doc, '查询模块的多种筛选条件均能正确工作。模糊搜索功能能在不输入完整关键词的情况下通过子串匹配找到目标记录。统计模块的 ASCII 柱状图直观展示数据分布。', sa=6)

    doc.add_page_break()

    # ── 第 4 章 ──────────────────────────────────────
    ah(doc, '4 课题收获与总结', 1)

    ah(doc, '4.1 理论知识应用体会', 2)
    ap(doc, '通过本次课程设计，我对 C/C++ 语言的链表、结构体、文件读写和字符串处理有了更深入的理解。与课堂上单独写算法不同，完整系统的开发需要把数据结构设计、输入校验、业务流程组织和异常处理整合在一起，任何一个环节的疏忽都可能导致系统行为不稳定甚至崩溃。', sa=6)
    ap(doc, '尤其是在租车和退车的状态联动设计中，我更加深刻地体会到数据一致性的重要性。一个看似简单的状态切换，实际上涉及 Vehicle、Renter 和 RentRecord 三个结构体的同步更新——车辆状态从可租变为已租再变回可租，用户的租车次数递增，租车记录从租用中变为已归还。这种多实体的联动设计思维，对后续学习数据库和业务系统开发具有重要的参考价值。', sa=6)
    ap(doc, '同时，项目的规范和注释要求在编码过程中帮助我养成了良好的开发习惯。函数命名采用动宾结构（如 addVehicle、findRenter、calcDateDiff），常量使用 constexpr，字符串转换使用 strtol/strtod，这些细节显著提高了代码的可维护性和可读性。', sa=6)

    ah(doc, '4.2 课题应用软件环境的体会', 2)
    ap(doc, '本项目采用命令行环境完成开发和测试。在开发过程中，CMake 和 Make 的使用帮助我理解了工程化构建的意义——随着源文件逐步拆分为核心文件加七个菜单模块，项目结构变得更加清晰，后期维护和调试也方便了很多。', sa=6)
    ap(doc, '在测试阶段深入终端通过反复操作验证系统功能，我逐渐理解了输入输出缓冲、文件指针管理、字符串截断这些 C 语言底层细节的重要性。例如，在开发初期曾因 fgets 保留换行符导致密码校验失败，后来通过 trimNewLine 统一处理所有字符串输入的末尾换行符解决了这个问题。', sa=6)

    ah(doc, '4.3 其它体会', 2)
    ap(doc, '在实际编码过程中，我意识到细节远比想象中的重要。例如输入校验的三种情况（空输入、非数字、超范围）、链表删除时的前驱指针维护、文件打开失败的检查等等，这些看似琐碎的点如果不处理好，就会直接影响用户体验甚至导致系统崩溃。课程设计让我从"能跑"阶段迈向了"能稳定跑"的阶段。', sa=6)
    ap(doc, '此外，本次项目让我对"设计与实现分离"有了初步认识。在编码之前需要先理清数据结构、函数接口和模块调用关系，否则后期修改成本会非常大。通过在纸上先画出结构体的字段、链表的关系图，再逐一实现各个功能，整个开发过程比最初预期的要顺利很多。', sa=6)
    ap(doc, '总体来说，汽车租赁管理系统虽然规模不大，但涵盖了一个管理信息系统应具备的全部要素：数据存储、业务流程、结果展示、统计分析和持久化读写。完成本项目后，我对如何将分散的知识点整合为可运行的系统有了更实际的认识，为后续的软件工程课程学习奠定了坚实基础。', sa=6)

    doc.add_page_break()

    # ── 第 5 章 ──────────────────────────────────────
    ah(doc, '5 参考文献', 1)
    for ref in [
        '[1] 谭浩强. C程序设计(第五版). 北京: 清华大学出版社, 2017.',
        '[2] Stanley B. Lippman et al. C++ Primer (5th Edition). Addison-Wesley, 2012.',
        '[3] 学术工程实践I指导书(汽车租赁管理系统-孙晶老师). 北方工业大学, 2026.',
        '[4] Brian W. Kernighan, Dennis M. Ritchie. The C Programming Language (2nd Ed). Prentice Hall, 1988.',
        '[5] Bjarne Stroustrup. The C++ Programming Language (4th Edition). Addison-Wesley, 2013.',
    ]:
        ap(doc, ref, fi=0, align=WD_ALIGN_PARAGRAPH.LEFT, sa=6)

    return doc


def main():
    doc = build()
    doc.save(OUT_PATH)
    print('报告已生成：' + OUT_PATH)


if __name__ == '__main__':
    main()
